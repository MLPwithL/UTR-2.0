from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\文件管理\东吴证券\UTR股票复现")
OUT = ROOT / "agent_workspace" / "outputs" / "summary.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "666666"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, name="Microsoft YaHei", size=10.5, color="000000", bold=False, italic=False):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])
    set_run_font(run, size=9, color=MID_GRAY)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.5 if level == 0 else 0.75)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.167
    set_run_font(p.add_run(text))
    return p


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_lead and text.startswith(bold_lead):
        set_run_font(p.add_run(bold_lead), bold=True)
        set_run_font(p.add_run(text[len(bold_lead):]))
    else:
        set_run_font(p.add_run(text))
    return p


def add_callout(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_geometry(table, [9360], 120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.10
    set_run_font(p.add_run(f"{label}："), color=DARK_BLUE, bold=True)
    set_run_font(p.add_run(text), color="243447")
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Bullet 2"):
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)


def build_document():
    doc = Document()
    configure_styles(doc)
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.add_run("东吴证券 UTR 2.0 因子复现项目｜周报"), size=8.5, color=MID_GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(footer.add_run("第 "), size=9, color=MID_GRAY)
    add_field(footer, "PAGE")
    set_run_font(footer.add_run(" 页"), size=9, color=MID_GRAY)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run("上周工作总结"), size=25, color="17365D", bold=True)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(14)
    set_run_font(p2.add_run("《技术分析十二——优加换手率 2.0》复现与口径审计"), size=13, color=MID_GRAY)

    meta = [
        ("汇报周期", "2026年6月15日—6月19日"),
        ("项目", "东吴证券 UTR 股票因子复现"),
        ("工作重点", "复现落地、流程对齐、交易约束与数据质量审计"),
        ("当前状态", "初版复现完成；关键口径已完成专项核查与小步改造"),
    ]
    for label, value in meta:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        set_run_font(p.add_run(f"{label}："), bold=True)
        set_run_font(p.add_run(value))

    add_callout(
        doc,
        "本周一句话总结",
        "完成 UTR 2.0 五因子复现框架和初版绩效报告，并围绕股票池、滚动窗口、交易执行与涨跌停约束开展二次审计，使项目从“结果复现”进一步推进到“口径可解释、流程可验证”。",
    )

    doc.add_heading("一、本周工作概述", level=1)
    add_body(
        doc,
        "上周围绕东吴证券《技术分析十二——优加换手率 2.0》开展完整复现。工作不止停留在因子公式和回测结果层面，还对两套实现流程进行了逐项比对，并针对 ST、停牌/无效交易、价格日期中断、涨跌停成交限制和持仓延续等关键环节进行了专项排查与改造。",
    )
    add_body(
        doc,
        "最终形成了因子计算、回测、报告输出和测试验证相互衔接的工程化流程，同时明确了当前结果与另一套测试代码之间仍存在的口径差异，为后续全量重跑和最终定稿建立了清晰路径。",
    )

    doc.add_heading("二、主要工作与阶段成果", level=1)
    doc.add_heading("1. 完成 UTR 2.0 初版复现与五因子统一回测", level=2)
    add_body(doc, "搭建并梳理了从原始数据扫描、基础因子计算、中性化处理到十分组回测、IC/RankIC 统计和图表输出的完整链路。统一覆盖以下五个因子：")
    for item in ("Turn20_neutral（传统低换手因子）", "STR_neutral（换手稳定性因子）", "UTR1.0", "UTR2.0", "UTR2.0_pure（剥离行业与 Barra 风格后的纯净因子）"):
        add_bullet(doc, item)
    add_body(doc, "初版报告覆盖 2006年1月至2023年3月、共207个月。按当时的 T+1 收盘执行口径，核心绩效结果如下。")

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    headers = ["因子", "年化对冲收益", "信息比率", "最大回撤"]
    rows = [
        ("Turn20_neutral", "33.53%", "1.95", "19.50%"),
        ("STR_neutral", "39.35%", "2.75", "9.82%"),
        ("UTR1.0", "38.70%", "3.12", "8.24%"),
        ("UTR2.0", "42.09%", "3.19", "7.65%"),
        ("UTR2.0_pure", "18.10%", "2.24", "10.05%"),
    ]
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_GRAY)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(cell.paragraphs[0].add_run(text), size=9.5, bold=True)
    for values in rows:
        cells = table.add_row().cells
        for i, text in enumerate(values):
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_run_font(cells[i].paragraphs[0].add_run(text), size=9.5)
    set_table_geometry(table, [3420, 1980, 1800, 2160], 120)
    add_body(doc, "说明：以上为初版复现口径结果。后续对回测交易时点和成交约束进行了调整，正式对外引用前需基于最新代码重新全量计算。", bold_lead="说明：")

    doc.add_heading("2. 对 scripts 与测试代码开展全流程差异审计", level=2)
    add_body(doc, "逐项对比工程化 scripts 流程与五步测试代码，确认两者并非简单的正式版/测试版关系，而是两条侧重点不同的实现路径：前者强调 Windows 环境下的一键复现和报告自动化，后者强调 SOP 拆步验证和中间矩阵留痕。")
    add_body(doc, "审计识别出的主要差异包括：")
    for item in (
        "新股过滤：60个自然日与60个交易日口径不同；",
        "Turn20/STR 窗口：dropna 后按个股滚动20条记录，与固定交易日窗口、允许部分缺失的实现不同；",
        "极值与标准化：3×1.4826×MAD 并多次 Z-score，与3×MAD、保留中性化残差尺度不同；",
        "回测执行：T+1 收盘调仓，与下月月初开盘买入、月末收盘卖出的 SOP 不同；",
        "十分组划分：手工切片与 rank + qcut 在边界样本上可能产生差异。",
    ):
        add_bullet(doc, item)
    add_callout(doc, "核心判断", "当前最可能造成结果不一致的前三项是滚动窗口、新股过滤和因子尺度处理；因此后续应按这三个方向依次统一，而不是直接对最终收益曲线做表面校准。")

    doc.add_heading("3. 完成股票池过滤逻辑的小步改造", level=2)
    add_body(doc, "为提高因子月度横截面的可交易性和与测试流程的一致性，对 calc_factors.py 做了两项局部调整：")
    add_bullet(doc, "ST 过滤由“仅判断月末是否 ST”改为整月半数日期规则：当月 ST 天数大于非 ST 天数时剔除，二者相等时保留。")
    add_bullet(doc, "新增基于 price.csv 的月内正常交易天数过滤：开盘价、收盘价有效且成交量大于0的日期须超过当月一半，否则剔除该股票当月样本。")
    add_body(doc, "这两项改造主要解决月末单点判断过于粗糙，以及 dropna 后滚动窗口可能拼接停牌前后有效记录的问题。ST 过滤新增单元测试并通过；price 月内交易日过滤已补充测试文件，但按当时要求尚未执行。")

    doc.add_heading("4. 完成回测成交约束改造", level=2)
    add_body(doc, "将 backtest.py 的核心交易框架调整为更接近原报告 SOP 的方式：下月首个交易日按复权开盘价买入、下月最后一个交易日按复权收盘价卖出，并显式模拟成交限制。")
    for item in (
        "月初开盘涨停（相对信号月末收盘涨幅超过9.8%）的股票视为无法买入，当期跳过；",
        "月末收盘跌停（相对前一交易日收盘跌幅超过9.8%）的股票视为无法卖出，保留在原分组并延续至下一期；",
        "为每个因子增加 holding summary 输出，记录目标持仓、新买入、carry 和下期延续数量，便于逐月复核。",
    ):
        add_bullet(doc, item)
    add_body(doc, "针对“涨停买不进、跌停卖不出并 carry”构造了最小回归样本，标准库 unittest 运行通过。")

    doc.add_heading("5. 开展 price.csv 日期中断专项检查", level=2)
    add_body(doc, "以全市场8,635个交易日构造基准日历，对5,835只股票在各自首末日期之间的记录完整性进行只读扫描。结果显示：")
    for item in (
        "267只股票存在日期中断，占比4.58%；",
        "合计缺失65,400个交易日；",
        "个别股票缺失比例超过20%，说明直接 dropna 后按记录数 rolling 确实可能跨越长停牌或数据空档。",
    ):
        add_bullet(doc, item)
    add_body(doc, "该检查为新增月内可交易性过滤提供了数据证据，也说明后续更彻底的方案应考虑固定全市场交易日窗口，而非仅依赖个股有效记录数。")

    doc.add_heading("6. 完成 UTR2.0 差异专项分析", level=2)
    add_body(doc, "确认两套代码的 UTR2.0 公式一致，差异重点不在公式本身，而在合成前 Turn20 与 STR 的尺度处理。UTR1 主要依赖排名，对线性尺度变化不敏感；UTR2 使用连续数值和 softsign 交互项，对中性化残差是否再次 Z-score 十分敏感。")
    add_callout(doc, "分析结论", "若 Turn20、STR、UTR1 表现接近而 UTR2 差异明显，应优先核对两项输入的横截面标准差和交互项分布，而不是先修改 UTR2 公式。")

    doc.add_heading("三、本周形成的主要产出", level=1)
    for item in (
        "形成 UTR 2.0 五因子复现最终报告及运行逻辑、参数说明；",
        "完成 scripts 与测试代码的流程差异分析和当前对齐状态复核；",
        "完成 ST 半数日期过滤改造及单元测试；",
        "完成 price 月内正常交易天数过滤改造并补充测试；",
        "完成回测月初开盘/月末收盘、涨跌停限制与 carry 机制改造及回归测试；",
        "完成 price.csv 日期缺口统计和 UTR2.0 差异专项分析。",
    ):
        add_bullet(doc, item)

    doc.add_heading("四、风险认识与当前边界", level=1)
    add_body(doc, "本周工作的一个重要收获，是将“复现结果好看”与“复现口径可验证”区分开来。当前项目已具备较完整的工程框架，但仍需注意以下边界：")
    for item in (
        "初版绩效尚未基于最新 ST、月内交易日和严格成交约束代码全量重跑，不能直接视为最终定稿数据；",
        "新股过滤仍存在60自然日与60交易日的口径差异；",
        "基础因子 rolling 仍采用 dropna 后20条有效记录，尚未完全改成固定交易日窗口；",
        "price.csv 与停牌事件表的口径可能不完全一致，需决定最终采用哪一套作为主标准；",
        "UTR2.0 的输入尺度尚需按“对齐测试代码”或“遵循报告等量纲逻辑”确定最终方案。",
    ):
        add_bullet(doc, item)

    doc.add_heading("五、下周工作计划", level=1)
    for item in (
        "优先统一 Turn20/STR 的固定20交易日窗口与最小有效天数规则，并补充边界测试；",
        "确定并统一新股过滤口径，建议按60个交易日执行；",
        "明确 3MAD、1.4826 修正和中性化后二次 Z-score 的最终研究口径；",
        "统一十分组划分方式，并核对边界月份分组成员；",
        "在上述口径确定后全量重跑 calc_factors.py 与 backtest.py，刷新因子底表、持仓摘要、绩效指标和图表；",
        "将新结果与研报原文、测试代码和初版报告进行三方核对，形成最终可交付版本。",
    ):
        add_bullet(doc, item)

    doc.add_heading("六、总体评价", level=1)
    add_body(doc, "上周已完成从“搭建复现框架”到“识别关键偏差并逐项修正”的主要过渡。当前工作的价值不仅是得到 UTR 2.0 的回测曲线，更在于建立了一套可追溯、可测试、可继续收敛的复现流程。下一阶段应以统一口径和全量重跑为主，确保最终报告中的每一个关键数字都能由最新代码稳定复现。")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_document())
